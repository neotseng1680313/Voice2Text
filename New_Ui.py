# -*- coding: utf-8 -*-
import re
import sys
import os
import speech_recognition as sr
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog, QLabel, QPushButton, QTextEdit, QComboBox, QMainWindow, QApplication
from docx import Document
from docx.shared import Pt
from transformers import pipeline  # 用於摘要生成
from subprocess import check_output, CalledProcessError, PIPE
from concurrent.futures import ThreadPoolExecutor

class Signals(QtCore.QObject):
    update_text = QtCore.pyqtSignal(str)

class Ui_MainWindow(QMainWindow):
    def __init__(self, parent=None):
        super(Ui_MainWindow, self).__init__(parent)
        self.lang_code = 'en-US'  # 設置一個默認語言代碼
        self.setupUi(self)  # 正確傳遞實例

    def setupUi(self, MainWindow):
        MainWindow.setWindowTitle("我的應用程式")
        # 根據需要添加其他UI設置
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(1237, 866)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")

        # Labels
        self.label = QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(530, 10, 71, 31))
        self.label.setObjectName("label")

        self.label_2 = QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(230, 70, 111, 51))
        self.label_2.setObjectName("label_2")

        # Buttons
        self.startButton = QPushButton(self.centralwidget)
        self.startButton.setGeometry(QtCore.QRect(400, 120, 93, 28))
        self.startButton.setObjectName("startButton")

        self.stopButton = QPushButton(self.centralwidget)
        self.stopButton.setGeometry(QtCore.QRect(500, 120, 93, 28))
        self.stopButton.setObjectName("stopButton")
        self.stopButton.setEnabled(False)

        self.uploadButton = QPushButton(self.centralwidget)
        self.uploadButton.setGeometry(QtCore.QRect(600, 120, 193, 28))
        self.uploadButton.setObjectName("uploadButton")

        self.summaryButton = QtWidgets.QPushButton(MainWindow)
        self.summaryButton.setGeometry(QtCore.QRect(450, 400, 200, 50))
        self.summaryButton.setObjectName("summaryButton")
        self.summaryButton.clicked.connect(self.generate_summary)

        # Language selector
        self.languageComboBox = QtWidgets.QComboBox(self.centralwidget)
        self.languageComboBox.setGeometry(QtCore.QRect(600, 160, 120, 28))
        self.languageComboBox.addItems(["English", "Chinese"])
        self.languageComboBox.setObjectName("languageComboBox")
        # Connect the language selection combobox change event
        self.languageComboBox.currentIndexChanged.connect(self.update_language_code)
        
        # Transcript Text Box
        self.transcriptText = QTextEdit(self.centralwidget)
        self.transcriptText.setGeometry(QtCore.QRect(230, 200, 671, 150))
        self.transcriptText.setObjectName("transcriptText")

        # Summarizer initialization
        self.summarizer = pipeline("summarization")

        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 1237, 25))
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        MainWindow.setStatusBar(self.statusbar)

        self.signals = Signals()
        self.startButton.clicked.connect(self.start_recording)
        self.stopButton.clicked.connect(self.stop_recording)
        self.uploadButton.clicked.connect(self.upload_audio_file)
        self.signals.update_text.connect(self.transcriptText.append)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def update_language_code(self, index):
        language_map = {0: 'en-US', 1: 'zh-TW'}
        self.lang_code = language_map.get(index, 'en-US')
        print(f"語言設置為: {self.lang_code}")  # 調試輸出當前語言設置

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "語音識別系統"))
        self.label.setText(_translate("MainWindow", "語音轉文字"))
        self.label_2.setText(_translate("MainWindow", "請點擊開始錄音"))
        self.startButton.setText(_translate("MainWindow", "開始"))
        self.stopButton.setText(_translate("MainWindow", "停止"))
        self.uploadButton.setText(_translate("MainWindow", "上傳語音文件"))
        self.summaryButton.setText(_translate("MainWindow", "生成摘要"))
        self.transcriptText.setPlaceholderText(_translate("MainWindow", "錄音內容將顯示在這裡。"))

    def set_language(self, lang):
        if lang == 'Chinese':
            self.lang_code = 'zh-TW'
        elif lang == 'English':
            self.lang_code = 'en-US'
        else:
            self.lang_code = 'en-US'  # 預設為英文

    def start_recording(self):
        # Initialize real-time speech recognition
        self.recognizer = sr.Recognizer()
        self.microphone = sr.Microphone()
        with self.microphone as source:
            self.recognizer.adjust_for_ambient_noise(source)
        self.stop_listening = self.recognizer.listen_in_background(self.microphone, self.speech_recognition_callback)
        self.stopButton.setEnabled(True)
        self.startButton.setEnabled(False)

    def stop_recording(self):
        # Stop real-time speech recognition
        self.stop_listening(wait_for_stop=False)
        self.stopButton.setEnabled(False)
        self.startButton.setEnabled(True)

    def speech_recognition_callback(self, recognizer, audio):
        try:
            # 调用 Google Web Speech API 进行语音识别
            # 注意: 这里默认使用了免费的API，对于高负载的商业用途可能需要使用API密钥
            text = recognizer.recognize_google(audio, language=self.lang_code)
            self.signals.update_text.emit(text)
        except sr.UnknownValueError:
            self.signals.update_text.emit("无法理解语音")
        except sr.RequestError as e:
            self.signals.update_text.emit(f"服务错误：{e}")

    def upload_audio_file(self):
        # Function to upload and process audio file
        filename, _ = QFileDialog.getOpenFileName(None, "選擇音頻文件", "", "Audio files (*.mp3 *.wav *.m4a);;All files (*)")
        if filename:
            self.process_audio_file(filename)
    '''
    def generate_summary(self):
        current_text = self.transcriptText.toPlainText()
        current_language = self.lang_code  # 確保有一個屬性或方法獲取當前語言代碼
        try:
            # 傳遞文本和語言代碼
            summary = check_output(['python', 'LangChain.py', current_text, current_language], universal_newlines=True, encoding='utf-8')
            self.transcriptText.setText(current_text + "\n\n" + summary)  # 在界面保留原文並顯示摘要
            self.save_text_to_word(current_text, summary)  # 保存原文和摘要到Word
        except CalledProcessError as e:
            print("錯誤:", e.output)
        except UnicodeDecodeError as e:
            print("解碼錯誤:", e)'''

    def generate_summary(self):
        current_text = self.transcriptText.toPlainText()
        current_language = self.lang_code  # 確保有一個屬性或方法獲取當前語言代碼
        try:
            # 傳遞文本和語言代碼
            summary = check_output(['python', 'LangChain.py', current_text, current_language], universal_newlines=True, encoding='utf-8')
        
            # 將摘要分割成條列式
            summary_lines = []
            parts = re.split(r'[。.]|\s{3,}', summary)
        
            for line in parts:
                if line.strip():
                    summary_lines.append(line.strip())

            list_summary = []
            for line in summary_lines:
                sub_parts = re.split(r'\s{3,}', line)
                for part in sub_parts:
                    combined = ' '.join(part.split())
                    if combined:
                        list_summary.append(f"{len(list_summary) + 1}. {combined}")

            formatted_summary = '\n'.join(list_summary)
        
            # 在界面保留原文並顯示摘要
            self.transcriptText.setText(current_text + "\n\n摘要:\n" + summary + "\n\n條列式摘要:\n" + formatted_summary)
        
            # 保存原文和摘要到Word
            self.save_text_to_word(current_text, summary, formatted_summary)
        except CalledProcessError as e:
            print("錯誤:", e.output)
        except UnicodeDecodeError as e:
            print("解碼錯誤:", e)

    def process_audio_file(self, filename):
        r = sr.Recognizer()
        with sr.AudioFile(filename) as source:
            audio_data = r.record(source)
        try:
            # 使用 Google Web Speech API 进行语音识别
            text = r.recognize_google(audio_data, language=self.lang_code)
            self.transcriptText.setText(text)  # 显示识别结果
        except (sr.UnknownValueError, sr.RequestError) as e:
            print(f"处理音频文件出错: {e}")

    '''
    def save_text_to_word(self, text, summary=None):
    # 嘗試打開現有Word文檔，如果不存在則創建一個新的
        try:
            doc = Document('processed_text.docx')
        except FileNotFoundError:
            doc = Document()

        #添加原文內容
        doc.add_paragraph("原文:\n" + text)

        # 如果有摘要，則添加摘要內容
        if summary:
            doc.add_paragraph("\n" + summary)

         # 保存文檔
        doc.save('processed_text.docx')
    '''
    def save_text_to_word(self, text, summary=None, list_summary=None):
        # 嘗試打開現有Word文檔，如果不存在則創建一個新的
        try:
            doc = Document('processed_text.docx')
        except FileNotFoundError:
            doc = Document()

        # 添加原文內容
        doc.add_paragraph("原文:\n" + text)

        # 如果有摘要，則添加摘要內容
        if summary:
            doc.add_paragraph("\n摘要:\n" + summary)
    
        # 如果有條列式摘要，則添加條列式摘要內容
        if list_summary:
            doc.add_paragraph("\n條列式摘要:\n" + list_summary)

        # 保存文檔
        doc.save('processed_text.docx')


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
